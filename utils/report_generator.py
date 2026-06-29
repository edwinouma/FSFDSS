from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

# ==========================================================
# GENERATE WORD REPORT
# ==========================================================

def generate_report(

    filename,

    latest_round,

    forecast_year,

    forecast_season,

    monthly_reference,

    seasonal_reference,

    national_fcs,

    national_rcsi,

    national_hhs,

    forecast_results,

    overall_risk,

    dominant,

    top_driver,

    indicator_table,

    forecast_table,

    driver_table,

    summary,

):

    doc = Document()

    # ------------------------------------------------------
    # TITLE
    # ------------------------------------------------------

    title = doc.add_heading(

        "Food Security Forecasting and Decision Support System",

        level=0

    )

    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_heading(

        "Automated Food Security Situation and Forecast Report",

        level=1

    )

    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()

    institution = doc.add_paragraph()

    institution.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    institution.add_run(

        "Master's Research Project\n"

    ).italic = True

    institution.add_run(

        "Food Security Forecasting and Decision Support System (FSFDSS)\n"

    ).bold = True

    institution.add_run(

        "Somalia"

    )

    # ------------------------------------------------------
    # REPORTING ROUND
    # ------------------------------------------------------

    p = doc.add_paragraph()

    p.add_run("Reporting Round: ").bold = True

    p.add_run(str(latest_round))

    p = doc.add_paragraph()

    p.add_run("Generated On: ").bold = True

    p.add_run(

        datetime.now().strftime(

            "%d %B %Y %H:%M"

        )

    )

    # ------------------------------------------------------
    # EXECUTIVE SUMMARY
    # ------------------------------------------------------

    doc.add_heading(

        "1. Executive Summary",

        level=1

    )

    doc.add_paragraph(summary)

    # ------------------------------------------------------
    # NATIONAL INDICATORS
    # ------------------------------------------------------

    doc.add_heading(

        "2. Historical Food Security Situation",

        level=1

    )

    table = doc.add_table(

        rows=1,

        cols=len(indicator_table.columns)

    )

    table.style = "Table Grid"

    # Header
    for i, col in enumerate(indicator_table.columns):
        table.rows[0].cells[i].text = str(col)

    # Data
    for _, row in indicator_table.iterrows():

        cells = table.add_row().cells

        for i, value in enumerate(row):

            if isinstance(value, float):

                cells[i].text = f"{value:.2f}"

            else:

                cells[i].text = str(value)

    # ------------------------------------------------------
    # MACHINE LEARNING FORECAST
    # ------------------------------------------------------

    doc.add_heading(

        "3. Machine Learning Forecast",

        level=1

    )

    doc.add_paragraph(

        f"""
    The Food Security Forecasting and Decision Support System
    (FSFDSS) generated baseline forecasts for the
    {forecast_season} {forecast_year} season using the latest
    available monthly monitoring data together with seasonal
    livelihood information and trained XGBoost machine learning
    models.

    National Forecast Results

    • Forecast Food Consumption Score (FCS):
      {forecast_results['national_fcs']:.2f}

    • Forecast Reduced Coping Strategy Index (rCSI):
      {forecast_results['national_rcsi']:.2f}

    • Forecast Household Hunger Scale (HHS):
      {forecast_results['national_hhs']:.2f}

    Overall National Risk Assessment

    {overall_risk.upper()}
    """

    )

    forecast_tbl = doc.add_table(

        rows=1,

        cols=len(forecast_table.columns)

    )

    forecast_tbl.style = "Table Grid"

    for i, col in enumerate(forecast_table.columns):
        forecast_tbl.rows[0].cells[i].text = str(col)

    for _, row in forecast_table.iterrows():

        cells = forecast_tbl.add_row().cells

        for i, value in enumerate(row):

            if isinstance(value, float):

                cells[i].text = f"{value:.2f}"

            else:

                cells[i].text = str(value)

    # ------------------------------------------------------
    # KEY DRIVERS
    # ------------------------------------------------------

    doc.add_heading(

        "4. Key Drivers of Food Security",

        level=1

    )

    doc.add_paragraph(

        f"""

Machine learning interpretation using SHAP values indicates that
{dominant.lower()} conditions remain the most influential drivers
of predicted food security outcomes across Somalia.

The single most influential predictor identified by the forecasting
models is:

• {top_driver}

These findings reinforce the importance of continuous monitoring of
environmental conditions, market dynamics, livelihood resilience,
and localized conflict when interpreting future food security risks.

"""

    )

    doc.add_paragraph()

    doc.add_heading(

        "Top Machine Learning Drivers",

        level=2

    )

    driver_tbl = doc.add_table(

        rows=1,

        cols=len(driver_table.columns)

    )

    driver_tbl.style = "Table Grid"

    # Header
    for i, col in enumerate(driver_table.columns):
        driver_tbl.rows[0].cells[i].text = str(col)

    # Rows
    for _, row in driver_table.iterrows():

        cells = driver_tbl.add_row().cells

        for i, value in enumerate(row):

            if isinstance(value, float):

                cells[i].text = f"{value:.3f}"

            else:

                cells[i].text = str(value)

    # ------------------------------------------------------
    # EARLY WARNING PRIORITIES
    # ------------------------------------------------------

    doc.add_heading(

        "5. Early Warning Priorities",

        level=1

    )

    priorities = [

        "Monitor rainfall performance and vegetation conditions (NDVI).",

        "Monitor staple food prices and Cost of the Minimum Basket (CMB).",

        "Track livestock market conditions, particularly goat prices.",

        "Monitor conflict incidents and conflict-related fatalities.",

        "Assess household livelihood resilience including debt levels, herd size and crop production."

    ]

    for priority in priorities:

        doc.add_paragraph(

            priority,

            style="List Bullet"

        )

    # ------------------------------------------------------
    # RECOMMENDATIONS
    # ------------------------------------------------------

    doc.add_heading(

        "6. Recommendations",

        level=1

    )

    recommendations = [

        "Continue seasonal monitoring of food security indicators.",

        "Strengthen environmental monitoring to support early warning.",

        "Increase surveillance of market prices and household purchasing power.",

        "Maintain routine monitoring of conflict trends in vulnerable livelihood zones.",

        "Update machine learning forecasts whenever new monthly data become available.",

        "Use the forecasting results alongside IPC analytical processes to support evidence-based decision-making."

    ]

    for recommendation in recommendations:

        doc.add_paragraph(

            recommendation,

            style="List Number"

        )

    # ------------------------------------------------------
    # FORECAST METADATA
    # ------------------------------------------------------

    doc.add_heading(

        "7. Forecast Metadata",

        level=1

    )

    doc.add_paragraph(
        f"Forecast Period: {forecast_season} {forecast_year}"
    )

    doc.add_paragraph(
        f"Historical Monthly Data Used: {monthly_reference}"
    )

    doc.add_paragraph(
        f"Seasonal Livelihood Information Used: {seasonal_reference}"
    )

    doc.add_paragraph(
        "Machine Learning Model: XGBoost Regression"
    )

    doc.add_paragraph(
        "Platform: Food Security Forecasting and Decision Support System (FSFDSS)"
    )

    doc.add_paragraph(
        f"Report Generated: {datetime.now():%d %B %Y %H:%M}"
    )

    # ------------------------------------------------------
    # CONCLUSION
    # ------------------------------------------------------

    doc.add_heading(

        "8. Conclusion",

        level=1

    )

    doc.add_paragraph(

        """
The Food Security Forecasting and Decision Support System (FSFDSS)
provides an integrated framework for combining historical monitoring,
machine learning forecasting and model interpretation to support
food security decision-making.

The system is intended to complement existing analytical approaches
by providing rapid evidence-based insights that can strengthen
preparedness, early warning and response planning.

Forecasts generated by the platform should be interpreted as
decision-support information and considered alongside expert
judgement and other operational evidence.

"""
    )

    # ------------------------------------------------------
    # SAVE REPORT
    # ------------------------------------------------------

    doc.save(filename)

    return filename

